from docx import Document
from docx.shared import Pt
import sys

MD_PATH = 'docs/components_doc.md'
DOCX_PATH = 'docs/NyayaVaani_UI_UX_Components.docx'

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 + max(0, (3-level)))


def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)


def main():
    doc = Document()
    doc.styles['Normal'].font.name = 'Inter'
    try:
        with open(MD_PATH, 'r', encoding='utf-8') as f:
            lines = [l.rstrip('\n') for l in f]
    except Exception as e:
        print('Failed to read markdown:', e)
        sys.exit(1)

    for line in lines:
        if line.startswith('=') or line.strip() == '':
            # ignore underline markers
            continue
        if line.startswith('#'):
            level = line.count('#', 0, 6)
            text = line.lstrip('#').strip()
            add_heading(doc, text, level=level)
        elif line.startswith('- ') or line.startswith('* '):
            # simple bullet
            p = doc.add_paragraph(line[2:].strip(), style='List Bullet')
        else:
            add_paragraph(doc, line)

    doc.save(DOCX_PATH)
    print('Saved DOCX to', DOCX_PATH)

if __name__ == '__main__':
    main()
